from __future__ import annotations

import os
import uuid
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document
from flask import Flask, abort, render_template, request, send_from_directory, url_for
from openpyxl import Workbook, load_workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from werkzeug.utils import secure_filename

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
ALLOWED_EXCEL = {"xlsx", "xlsm"}
ALLOWED_WORD = {"docx"}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024


def extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def make_name(stem: str, suffix: str) -> str:
    safe_stem = secure_filename(stem) or "converted_file"
    return f"{safe_stem}_{uuid.uuid4().hex[:8]}.{suffix}"


def excel_to_pdf(source: Path, destination: Path) -> None:
    workbook = load_workbook(source, data_only=True, read_only=True)
    styles = getSampleStyleSheet()
    story = []

    for sheet_number, sheet in enumerate(workbook.worksheets):
        rows = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                rows.append(values)

        story.append(Paragraph(sheet.title, styles["Heading2"]))
        story.append(Spacer(1, 4 * mm))
        if not rows:
            story.append(Paragraph("This worksheet is empty.", styles["BodyText"]))
        else:
            column_count = max(len(row) for row in rows)
            normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
            available_width = landscape(A4)[0] - 24 * mm
            col_width = available_width / min(column_count, 10)
            table = Table(normalized_rows, repeatRows=1, colWidths=[col_width] * column_count)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f766e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(table)
        if sheet_number < len(workbook.worksheets) - 1:
            story.append(Spacer(1, 12 * mm))

    SimpleDocTemplate(
        str(destination), pagesize=landscape(A4),
        rightMargin=12 * mm, leftMargin=12 * mm, topMargin=14 * mm, bottomMargin=14 * mm,
        title=source.stem,
    ).build(story)
    workbook.close()


def pdf_to_excel(source: Path, destination: Path) -> None:
    workbook = Workbook()
    workbook.remove(workbook.active)
    has_content = False

    with pdfplumber.open(source) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            worksheet = workbook.create_sheet(f"Page {index}")
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        worksheet.append(["" if cell is None else cell for cell in row])
                    worksheet.append([])
                    has_content = True
            else:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    worksheet.append([line])
                    has_content = True
            worksheet.freeze_panes = "A2"
            worksheet.column_dimensions["A"].width = 40

    if not has_content:
        workbook.create_sheet("Extracted content").append(["No extractable text or tables were found."])
    workbook.save(destination)


def word_to_excel(source: Path, destination: Path) -> None:
    document = Document(source)
    workbook = Workbook()
    text_sheet = workbook.active
    text_sheet.title = "Document text"
    text_sheet.append(["Paragraph"])
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_sheet.append([paragraph.text.strip()])
    text_sheet.column_dimensions["A"].width = 90
    text_sheet.freeze_panes = "A2"

    for index, table in enumerate(document.tables, start=1):
        worksheet = workbook.create_sheet(f"Table {index}")
        for row in table.rows:
            worksheet.append([cell.text.strip() for cell in row.cells])
        worksheet.freeze_panes = "A2"
    workbook.save(destination)


def word_to_pdf(source: Path, destination: Path) -> None:
    document = Document(source)
    styles = getSampleStyleSheet()
    story = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            style = styles["Heading2"] if paragraph.style and paragraph.style.name.startswith("Heading") else styles["BodyText"]
            story.extend([Paragraph(escaped, style), Spacer(1, 3 * mm)])
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            width = (A4[0] - 30 * mm) / max(len(row) for row in rows)
            pdf_table = Table(rows, repeatRows=1, colWidths=[width] * max(len(row) for row in rows))
            pdf_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f766e")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9), ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
            story.extend([pdf_table, Spacer(1, 5 * mm)])
    if not story:
        story.append(Paragraph("This Word document is empty.", styles["BodyText"]))
    SimpleDocTemplate(str(destination), pagesize=A4, rightMargin=15 * mm, leftMargin=15 * mm, topMargin=16 * mm, bottomMargin=16 * mm, title=source.stem).build(story)


def convert_file(source: Path, mode: str, destination: Path) -> None:
    converters = {
        "excel-to-pdf": excel_to_pdf,
        "pdf-to-excel": pdf_to_excel,
        "word-to-excel": word_to_excel,
        "word-to-pdf": word_to_pdf,
    }
    converters[mode](source, destination)


def create_zip(outputs: list[dict[str, str]]) -> str:
    zip_name = f"sheetshift_batch_{uuid.uuid4().hex[:8]}.zip"
    with zipfile.ZipFile(OUTPUT_DIR / zip_name, "w", zipfile.ZIP_DEFLATED) as archive:
        for output in outputs:
            archive.write(OUTPUT_DIR / output["name"], arcname=output["download_name"])
    return zip_name


@app.route("/", methods=["GET", "POST"])
def index():
    result = None
    error = None
    selected_mode = request.form.get("mode", "excel-to-pdf")

    if request.method == "POST":
        uploaded_files = [item for item in request.files.getlist("file") if item.filename]
        if not uploaded_files:
            error = "Choose a file before converting."
        else:
            expected_by_mode = {"excel-to-pdf": ALLOWED_EXCEL, "pdf-to-excel": {"pdf"}, "word-to-excel": ALLOWED_WORD, "word-to-pdf": ALLOWED_WORD}
            expected = expected_by_mode.get(selected_mode, set())
            invalid_file = next((item for item in uploaded_files if extension(item.filename) not in expected), None)
            if invalid_file:
                error = {"excel-to-pdf": "Please upload Excel files only (.xlsx or .xlsm).", "pdf-to-excel": "Please upload PDF files only.", "word-to-excel": "Please upload Word files only (.docx).", "word-to-pdf": "Please upload Word files only (.docx)."}.get(selected_mode, "Choose a valid conversion type.")
            else:
                converted_outputs = []
                failed_files = []
                output_ext = "pdf" if selected_mode in {"excel-to-pdf", "word-to-pdf"} else "xlsx"
                for uploaded_file in uploaded_files:
                    file_ext = extension(uploaded_file.filename)
                    source_name = make_name(Path(uploaded_file.filename).stem, file_ext)
                    source_path = UPLOAD_DIR / source_name
                    output_name = make_name(Path(uploaded_file.filename).stem, output_ext)
                    output_path = OUTPUT_DIR / output_name
                    uploaded_file.save(source_path)
                    try:
                        convert_file(source_path, selected_mode, output_path)
                        converted_outputs.append({"name": output_name, "download_name": f"{Path(uploaded_file.filename).stem}.{output_ext}"})
                    except Exception:
                        app.logger.exception("Conversion failed for %s", uploaded_file.filename)
                        output_path.unlink(missing_ok=True)
                        failed_files.append(uploaded_file.filename)
                    finally:
                        source_path.unlink(missing_ok=True)
                if converted_outputs:
                    result = {"outputs": converted_outputs, "label": "PDF" if output_ext == "pdf" else "Excel workbook", "zip_name": create_zip(converted_outputs) if len(converted_outputs) > 1 else None}
                if failed_files:
                    failure_message = f"Could not convert: {', '.join(failed_files)}."
                    error = failure_message if not converted_outputs else f"{failure_message} The other files are ready below."

    return render_template("index.html", result=result, error=error, selected_mode=selected_mode)


@app.route("/download/<path:filename>")
def download(filename: str):
    if Path(filename).name != filename or not (OUTPUT_DIR / filename).is_file():
        abort(404)
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True, download_name=filename)


@app.errorhandler(413)
def too_large(_error):
    return render_template("index.html", error="Files must be 16 MB or smaller.", result=None, selected_mode="excel-to-pdf"), 413


if __name__ == "__main__":
    UPLOAD_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)
    app.run(debug=True)
